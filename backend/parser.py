"""
backend/parser.py
Parses PDF, DOCX, and TXT documents into raw text.
"""
from __future__ import annotations

import io
from pathlib import Path


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF (fitz)."""
    try:
        import fitz  # pymupdf

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = [page.get_text("text") for page in doc]
        doc.close()
        text = "\n".join(pages).strip()
        if not text:
            raise ValueError("No extractable text found in PDF (may be image-based).")
        return text
    except ImportError:
        # Fallback to pdfplumber if pymupdf unavailable
        import pdfplumber

        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        text = "\n".join(text_parts).strip()
        if not text:
            raise ValueError("No extractable text found in PDF.")
        return text


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError("No extractable text found in DOCX.")
    return text


def parse_txt(file_bytes: bytes) -> str:
    """Decode plain-text bytes."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode TXT file with common encodings.")


def parse_document(filename: str, file_bytes: bytes) -> str:
    """
    Dispatch to the correct parser based on file extension.

    Args:
        filename: Original filename (used to detect extension).
        file_bytes: Raw bytes of the uploaded file.

    Returns:
        Extracted plain text.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_bytes)
    elif ext == ".txt":
        return parse_txt(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            "Please upload PDF, DOCX, or TXT files."
        )
